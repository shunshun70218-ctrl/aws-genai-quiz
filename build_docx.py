#!/usr/bin/env python3
"""Regenerate the corrected Word document from questions_full.json.
Requires: pip install python-docx
Usage:    python3 build_docx.py
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor

qs = json.load(open('questions_full.json', encoding='utf-8'))
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

for q in qs:
    h = doc.add_paragraph(); r = h.add_run(f"Question #{q['num']}")
    r.bold = True; r.font.size = Pt(13)
    doc.add_paragraph(q['stem'])
    for L in sorted(q['options']):
        doc.add_paragraph(f"{L}. {q['options'][L]}")
    doc.add_paragraph("")
    ans = doc.add_paragraph(); ar = ans.add_run(f"answer: {', '.join(q['answers'])}")
    ar.bold = True; ar.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    ep = doc.add_paragraph(); er = ep.add_run("考點背法 \U0001F3AF"); er.bold = True
    if q['explain']:
        doc.add_paragraph(q['explain'])

doc.save('AWS_GenAI_Q1-Q97_題目格式_修正版.docx')
print("saved AWS_GenAI_Q1-Q97_題目格式_修正版.docx")
